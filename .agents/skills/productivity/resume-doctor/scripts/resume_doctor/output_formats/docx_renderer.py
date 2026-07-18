"""
DOCX Renderer — Microsoft Word format via python-docx.

Generates professional .docx files compatible with Word, Google Docs, LibreOffice.
"""
from pathlib import Path
from typing import List
from io import BytesIO

from resume_doctor.models.resume_data import ResumeData, NDALevel
from resume_doctor.output_formats import BaseRenderer, OutputFormat


class DocxRenderer(BaseRenderer):
    """Render ResumeData to DOCX using python-docx."""

    format = OutputFormat.DOCX
    extension = ".docx"

    def _apply_nda(self, text: str, level: NDALevel) -> str:
        if level == NDALevel.L0:
            return text
        import re
        if level == NDALevel.L1:
            return re.sub(r'\b[A-Z][a-z]+ (?:Inc|Corp|LLC|Ltd|Technologies|Systems|Labs)\b', '[Company]', text)
        elif level == NDALevel.L2:
            return re.sub(r'\b(?:Senior|Lead|Principal|Staff)\s+\w+', '[Role]', text)
        return text

    def render(self, resume: ResumeData) -> bytes:
        """Render to DOCX bytes."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            from docx.oxml.ns import qn
        except ImportError:
            raise RuntimeError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        doc = Document()

        # Configure default style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.15

        # Configure heading styles
        for level in range(1, 4):
            h_style = doc.styles[f'Heading {level}']
            h_style.font.name = 'Calibri'
            h_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            if level == 1:
                h_style.font.size = Pt(14)
                h_style.font.bold = True
            elif level == 2:
                h_style.font.size = Pt(12)
                h_style.font.bold = True
            elif level == 3:
                h_style.font.size = Pt(11)
                h_style.font.bold = True

        # Adjust margins
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # --- CONTACT HEADER ---
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run(resume.contact.name)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        if resume.contact.headline:
            headline_para = doc.add_paragraph()
            headline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = headline_para.add_run(resume.contact.headline)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)

        # Contact info line
        contact_parts = []
        if resume.contact.location:
            contact_parts.append(resume.contact.location)
        if resume.contact.phone:
            contact_parts.append(resume.contact.phone)
        if resume.contact.email:
            contact_parts.append(resume.contact.email)
        for link in resume.contact.links[:3]:
            if link.url:
                contact_parts.append(link.label or link.url)

        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact_para.add_run(" | ".join(contact_parts))
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Horizontal line
        line_para = doc.add_paragraph()
        line_para.paragraph_format.space_before = Pt(4)
        line_para.paragraph_format.space_after = Pt(8)
        run = line_para.add_run("━" * 80)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # --- PROFESSIONAL SUMMARY ---
        if resume.summary:
            doc.add_heading('Professional Summary', level=1)
            p = doc.add_paragraph(resume.summary)
            p.paragraph_format.space_after = Pt(8)

        # --- SKILLS ---
        if resume.skills.categories or resume.skills.summary:
            doc.add_heading('Skills & Technical Proficiency', level=1)

            if resume.skills.summary:
                p = doc.add_paragraph(resume.skills.summary)
                p.paragraph_format.space_after = Pt(6)

            if resume.skills.categories:
                for cat in resume.skills.categories:
                    if cat.skills:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(f"{cat.name}: ")
                        run.bold = True
                        skill_names = [s.name for s in cat.skills]
                        p.add_run(", ".join(skill_names))
                        p.paragraph_format.space_after = Pt(2)

        # --- EXPERIENCE ---
        if resume.experience:
            doc.add_heading('Experience', level=1)

            for entry in resume.experience:
                # Header: Role | Company | Location | Dates
                header_para = doc.add_paragraph()
                header_para.paragraph_format.space_before = Pt(8)
                header_para.paragraph_format.space_after = Pt(2)

                parts = [
                    (entry.role, True),
                    (entry.company, True),
                ]
                if entry.location:
                    parts.append((entry.location, False))
                parts.append((entry.date_range, False))

                for i, (text, bold) in enumerate(parts):
                    run = header_para.add_run(text)
                    run.bold = bold
                    run.font.size = Pt(10.5)
                    if i < len(parts) - 1:
                        header_para.add_run(" | ").font.color.rgb = RGBColor(0x66, 0x66, 0x66)

                if entry.description:
                    p = doc.add_paragraph(entry.description)
                    p.paragraph_format.space_after = Pt(4)

                if entry.bullets:
                    for bullet in entry.bullets:
                        bullet_text = self._apply_nda(bullet, entry.nda_level)
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(bullet_text)
                        p.paragraph_format.left_indent = Inches(0.3)
                        p.paragraph_format.space_after = Pt(2)

                if entry.technologies:
                    p = doc.add_paragraph()
                    run = p.add_run("Technologies: ")
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.italic = True
                    p.add_run(", ".join(entry.technologies)).font.size = Pt(9.5)
                    p.paragraph_format.space_after = Pt(2)

                if entry.metrics:
                    p = doc.add_paragraph()
                    run = p.add_run("Key Metrics: ")
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.italic = True
                    p.add_run("; ".join(entry.metrics)).font.size = Pt(9.5)
                    p.paragraph_format.space_after = Pt(2)

        # --- PROJECTS ---
        if resume.projects:
            doc.add_heading('Selected Projects', level=1)

            for proj in resume.projects:
                header_para = doc.add_paragraph()
                header_para.paragraph_format.space_before = Pt(8)
                header_para.paragraph_format.space_after = Pt(2)

                run = header_para.add_run(proj.name)
                run.bold = True
                run.font.size = Pt(11)
                if proj.role:
                    header_para.add_run(" | ").font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    run = header_para.add_run(proj.role)
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

                if proj.description:
                    p = doc.add_paragraph(proj.description)
                    p.paragraph_format.space_after = Pt(4)

                if proj.bullets:
                    for bullet in proj.bullets:
                        bullet_text = self._apply_nda(bullet, proj.nda_level)
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(bullet_text)
                        p.paragraph_format.left_indent = Inches(0.3)
                        p.paragraph_format.space_after = Pt(2)

                if proj.technologies:
                    p = doc.add_paragraph()
                    run = p.add_run("Stack: ")
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.italic = True
                    p.add_run(", ".join(proj.technologies)).font.size = Pt(9.5)
                    p.paragraph_format.space_after = Pt(2)

                if proj.link:
                    p = doc.add_paragraph()
                    run = p.add_run("Link: ")
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run = p.add_run(proj.link)
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
                    run.font.underline = True
                    p.paragraph_format.space_after = Pt(2)

        # --- EDUCATION ---
        if resume.education:
            doc.add_heading('Education & Certifications', level=1)

            for edu in resume.education:
                header_para = doc.add_paragraph()
                header_para.paragraph_format.space_before = Pt(6)
                header_para.paragraph_format.space_after = Pt(2)

                parts = [
                    (edu.degree, True),
                    (edu.institution, False),
                ]
                if edu.location:
                    parts.append((edu.location, False))
                if edu.graduation_date:
                    parts.append((edu.graduation_date, False))

                for i, (text, bold) in enumerate(parts):
                    run = header_para.add_run(text)
                    run.bold = bold
                    run.font.size = Pt(10.5)
                    if i < len(parts) - 1:
                        header_para.add_run(" | ").font.color.rgb = RGBColor(0x66, 0x66, 0x66)

                if edu.honors:
                    p = doc.add_paragraph()
                    run = p.add_run("Honors: ")
                    run.bold = True
                    run.font.size = Pt(9.5)
                    p.add_run(", ".join(edu.honors)).font.size = Pt(9.5)
                    p.paragraph_format.space_after = Pt(2)

                if edu.relevant_coursework:
                    p = doc.add_paragraph()
                    run = p.add_run("Relevant Coursework: ")
                    run.bold = True
                    run.font.size = Pt(9.5)
                    p.add_run(", ".join(edu.relevant_coursework)).font.size = Pt(9.5)
                    p.paragraph_format.space_after = Pt(2)

        # --- CERTIFICATIONS ---
        if resume.certifications:
            doc.add_heading('Certifications', level=1)

            for cert in resume.certifications:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)

                parts = [
                    (cert.name, True),
                    (cert.issuer, False),
                ]
                if cert.date_earned:
                    parts.append((cert.date_earned, False))

                for i, (text, bold) in enumerate(parts):
                    run = p.add_run(text)
                    run.bold = bold
                    run.font.size = Pt(10)
                    if i < len(parts) - 1:
                        p.add_run(" | ").font.color.rgb = RGBColor(0x66, 0x66, 0x66)

                if cert.credential_url:
                    p.add_run(" | ")
                    run = p.add_run("Verify")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
                    run.font.underline = True

        # Save to bytes
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.read()

    def save(self, resume: ResumeData, output_path: Path) -> Path:
        """Render and save to .docx file."""
        content = self.render(resume)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(content)
        return output_path

    def get_instructions(self, resume: ResumeData) -> str:
        name = resume.contact.name or "Your Name"
        return f"""# How to Use Your DOCX Resume

## Quick Start
The generated `.docx` file opens in:
- **Microsoft Word** (Windows/Mac)
- **Google Docs** (drag & drop or File → Open)
- **LibreOffice Writer** (free, cross-platform)
- **Apple Pages** (imports with good fidelity)

## Recruiter-Friendly
- Most ATS systems parse DOCX natively
- Recruiters expect Word format for editing/annotation
- Track Changes works for collaborative review

## Editing & Customization
1. **Open in Word/Google Docs**
2. **Edit directly** — all text is editable
3. **Style adjustments**:
   - Font: Calibri 10.5pt (professional default)
   - Margins: 0.6" top/bottom, 0.75" sides
   - Line spacing: 1.15
4. **Save as PDF** for final distribution: File → Save As → PDF

## Conversion
```bash
# Convert to PDF (LibreOffice headless)
libreoffice --headless --convert-to pdf resume.docx

# Convert to HTML (Pandoc)
pandoc resume.docx -o resume.html

# Convert to Markdown
pandoc resume.docx -o resume.md
```

## Google Docs Workflow
1. Upload to Google Drive
2. Right-click → Open with → Google Docs
3. Share link with "Viewer" or "Commenter" access
4. Collaborate with reviewers in real-time

## ATS Compatibility
- **Standard structure**: Headings, lists, tables parsed correctly
- **No text boxes/columns** that confuse parsers
- **Embedded hyperlinks** preserved (email, LinkedIn, portfolio)
- **Keywords** in natural context

## Companion Files
- `resume_data.json` — Canonical structured data (source of truth)
- `resume.md` — Markdown version (Git-friendly)
- `resume.html` — Web version
- `latex/main.tex` — Overleaf version
- `pdf/resume.pdf` — Print-ready PDF
- `validation-report.json` — Quality gates (all should PASS)

## Next Steps
- **Tailor for another job**: `resume-doctor tailor --resume resume_data.json --job <url> --format docx`
- **Export other formats**: `resume-doctor build --resume resume_data.json --format pdf|latex|markdown|html`
- **Print/Share**: Open in Word → File → Print → Save as PDF
"""


from resume_doctor.output_formats import RendererRegistry
RendererRegistry.register(OutputFormat.DOCX, DocxRenderer)